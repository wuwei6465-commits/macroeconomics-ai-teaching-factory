#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
题库工厂 - 学习通格式题库批量转换器

功能：将Markdown格式的习题批量转换为学习通快速导入格式
作者：吴伟
单位：广东白云学院
日期：2025年
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def parse_markdown_exercises(md_content: str) -> list:
    """
    解析Markdown格式的习题内容
    
    参数:
        md_content: Markdown文件内容
    
    返回:
        习题列表，每个习题为字典格式
    """
    exercises = []
    
    # 按题目分割（假设题目以"### 题目"或数字编号开头）
    # 这里需要根据实际格式调整正则表达式
    pattern = r'#{1,3}\s*(?:题目|习题)?\s*(\d+)[\.、]\s*\n'
    parts = re.split(pattern, md_content)
    
    for i in range(1, len(parts), 2):
        if i < len(parts):
            question_num = parts[i]
            question_content = parts[i + 1] if i + 1 < len(parts) else ""
            
            # 解析题目内容
            exercise = parse_single_exercise(question_num, question_content)
            if exercise:
                exercises.append(exercise)
    
    return exercises


def parse_single_exercise(num: str, content: str) -> dict:
    """
    解析单个习题
    
    参数:
        num: 题号
        content: 题目内容
    
    返回:
        习题字典，包含题干、选项、答案、解析等
    """
    exercise = {
        'num': num,
        'type': '单选题',  # 默认类型
        'question': '',
        'options': [],
        'answer': '',
        'analysis': '',
        'difficulty': '中',
        'bloom_level': 'L2'  # 布鲁姆认知能级
    }
    
    lines = content.strip().split('\n')
    
    # 解析题干（第一行）
    if lines:
        exercise['question'] = lines[0].strip()
    
    # 解析选项（A、B、C、D开头）
    option_pattern = r'^([A-D])[\.、\s]+(.+)$'
    for line in lines[1:]:
        match = re.match(option_pattern, line.strip())
        if match:
            option_letter = match.group(1)
            option_text = match.group(2)
            exercise['options'].append(f"{option_letter}. {option_text}")
    
    # 解析答案（【答案】或答案：开头）
    answer_pattern = r'(?:【答案】|答案[:：])\s*([A-D]+)'
    for line in lines:
        match = re.search(answer_pattern, line)
        if match:
            exercise['answer'] = match.group(1)
    
    # 解析解析（【解析】或解析：开头）
    analysis_pattern = r'(?:【解析】|解析[:：])\s*(.+)'
    for line in lines:
        match = re.search(analysis_pattern, line)
        if match:
            exercise['analysis'] = match.group(1)
    
    # 解析难度（【难度】或难度：开头）
    difficulty_pattern = r'(?:【难度】|难度[:：])\s*([易中难])'
    for line in lines:
        match = re.search(difficulty_pattern, line)
        if match:
            exercise['difficulty'] = match.group(1)
    
    # 解析布鲁姆能级（【L1-L6】或能级：开头）
    bloom_pattern = r'(?:【L(\d)】|能级[:：])\s*L(\d)'
    for line in lines:
        match = re.search(bloom_pattern, line)
        if match:
            exercise['bloom_level'] = f"L{match.group(1) or match.group(2)}"
    
    return exercise


def convert_to_learning_platform_format(exercises: list) -> str:
    """
    转换为学习通快速导入格式
    
    学习通格式示例：
    1. 题目内容
    A. 选项A
    B. 选项B
    C. 选项C
    D. 选项D
    答案：A
    解析：本题考查...
    
    参数:
        exercises: 习题列表
    
    返回:
        学习通格式的文本内容
    """
    output_lines = []
    
    for i, ex in enumerate(exercises, 1):
        # 题号+题干
        output_lines.append(f"{i}. {ex['question']}")
        
        # 选项
        for option in ex['options']:
            output_lines.append(option)
        
        # 答案
        output_lines.append(f"答案：{ex['answer']}")
        
        # 解析
        if ex['analysis']:
            output_lines.append(f"解析：{ex['analysis']}")
        
        # 难度和能级（可选，学习通可能不识别）
        output_lines.append(f"难度：{ex['difficulty']}  认知能级：{ex['bloom_level']}")
        
        # 空行分隔
        output_lines.append("")
    
    return '\n'.join(output_lines)


def batch_convert(
    input_dir: str,
    output_dir: str,
    output_format: str = 'word'  # 'word' 或 'txt'
) -> list:
    """
    批量转换Markdown习题为学习通格式
    
    参数:
        input_dir: 输入目录（包含.md文件）
        output_dir: 输出目录
        output_format: 输出格式，'word' 或 'txt'
    
    返回:
        生成的文件路径列表
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    generated_files = []
    
    # 遍历所有Markdown文件
    for md_file in sorted(input_path.glob('*.md')):
        print(f'处理文件: {md_file.name}')
        
        try:
            # 读取Markdown内容
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 解析习题
            exercises = parse_markdown_exercises(md_content)
            
            if not exercises:
                print(f'  警告: 未解析到习题 - {md_file.name}')
                continue
            
            # 生成输出
            chapter_name = md_file.stem  # 文件名（不含扩展名）
            
            if output_format == 'word':
                # 生成Word文档
                output_file = output_path / f'{chapter_name}_学习通题库.docx'
                generate_word_document(exercises, output_file, chapter_name)
            else:
                # 生成纯文本
                output_file = output_path / f'{chapter_name}_学习通题库.txt'
                content = convert_to_learning_platform_format(exercises)
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            generated_files.append(str(output_file))
            print(f'  成功: 生成 {len(exercises)} 道题目 - {output_file.name}')
            
        except Exception as e:
            print(f'  错误: 处理失败 - {e}')
            continue
    
    print(f'\n批量转换完成：成功 {len(generated_files)} 个文件')
    return generated_files


def generate_word_document(exercises: list, output_file: Path, chapter_name: str):
    """
    生成Word格式的题库文档
    
    参数:
        exercises: 习题列表
        output_file: 输出文件路径
        chapter_name: 章节名称
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(f'{chapter_name} 题库', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加说明
    doc.add_paragraph(f'共 {len(exercises)} 道题目')
    doc.add_paragraph()
    
    # 添加题目
    for i, ex in enumerate(exercises, 1):
        # 题号+题干
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(ex['question'])
        
        # 选项
        for option in ex['options']:
            doc.add_paragraph(option, style='List Bullet')
        
        # 答案
        p = doc.add_paragraph()
        p.add_run('答案：').bold = True
        p.add_run(ex['answer'])
        
        # 解析
        if ex['analysis']:
            p = doc.add_paragraph()
            p.add_run('解析：').bold = True
            p.add_run(ex['analysis'])
        
        # 难度和能级
        p = doc.add_paragraph()
        p.add_run(f'难度：{ex["difficulty"]}  认知能级：{ex["bloom_level"]}').italic = True
        
        # 空行分隔
        doc.add_paragraph()
    
    # 保存文档
    doc.save(output_file)


# 使用示例
if __name__ == '__main__':
    # 示例：批量转换
    batch_convert(
        input_dir='examples/chapter11-isl/exercises/',
        output_dir='output/题库/',
        output_format='word'
    )
